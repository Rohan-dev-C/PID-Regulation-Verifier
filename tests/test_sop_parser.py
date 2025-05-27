import os
from docx import Document
from src.sop_parser import SOPParser

def test_sop_parser_extracts_allcaps(tmp_path):
    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Initialize PUMP and VALVE before operation.")
    doc.add_paragraph("Ensure SENSOR is calibrated.")
    doc.save(str(p))

    parser = SOPParser(sop_path=p)
    result = parser.parse()

    assert "step_0" in result
    assert set(result["step_0"]) == {"PUMP", "VALVE"}

    assert "step_1" in result
    assert result["step_1"] == ["SENSOR"]
