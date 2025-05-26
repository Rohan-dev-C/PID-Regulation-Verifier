"""
Rich SOP parser – reads every part of the DOCX:
• Normal paragraphs
• Bullets / numbered lists
• Tables (each cell)
Extracts ALL-CAPS or tag-like tokens as component labels.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Dict, List

from docx import Document

from .config import settings
from .utils import sent_tokenize

LOGGER = logging.getLogger(__name__)

# Token regex: ALLCAPS or mixed like “PUMP-101”, “V1”, “T_SENSOR”
COMP_REGEX = re.compile(r"[A-Z][A-Z0-9_-]{1,}")


class SOPParser:
    def __init__(self, sop_path: Path | str | None = None) -> None:
        self.sop_path = Path(sop_path or settings.sop_path)

    # ---------------------------------------------------------------------
    def parse(self) -> Dict[str, List[str]]:
        """
        Returns { step_id : [components...] }.

        If the SOP .docx is missing, returns {} and logs a warning.
        """
        if not self.sop_path.is_file():
            LOGGER.warning("SOP file not found: %s (skipping SOP checks)", self.sop_path)
            return {}

        LOGGER.info("Parsing SOP %s", self.sop_path)
        doc = Document(str(self.sop_path))
        comp_by_step: Dict[str, List[str]] = {}

        step_idx = 0

        # ── 1) Paragraphs (incl. bullets / numbering) ────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            step_id = f"para_{step_idx}"
            comp_by_step[step_id] = self._extract_components(text)
            step_idx += 1

        # ── 2) Table cells ───────────────────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    step_id = f"table{t_idx}_r{r_idx}c{c_idx}"
                    comp_by_step[step_id] = self._extract_components(text)

        LOGGER.info("Extracted %d SOP steps (paragraphs + table cells)", len(comp_by_step))
        return comp_by_step

    # ---------------------------------------------------------------------
    @staticmethod
    def _extract_components(text: str) -> List[str]:
        """
        Return a list of component labels found in `text`
        using regex + sentence tokenisation.
        """
        labels: List[str] = []
        for sentence in sent_tokenize(text):
            labels.extend(COMP_REGEX.findall(sentence))
        return labels
